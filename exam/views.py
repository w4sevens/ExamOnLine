from django.shortcuts import render, HttpResponse, redirect
from django.http import JsonResponse
from django.views import View
from django.utils.safestring import mark_safe

import os, json
import datetime
import math

from exam import examforms, models
from users.models import Units
from exam.models import QuestionInfo, QuestionBank
from exam.models import TestPaper
from exam.examforms import QuestionInfoModelForm, TestPaperModelForm

from utils.myclass_word import QuestionBankClass
from utils.shuffle import shuffle
from utils.utils import dy, get_pagenator



# 考试中心-首页
def paper_info(request):

    # todo 用户后注册，则登录后有很多考试需要参加！！！

    # 显示正在进行的考试情况
    # papers = models.TestPaper.objects.filter(chance_times=1, add_time__gt=request.user.date_joined).all().order_by('-add_time')
    papers = models.TestPaper.objects.filter(chance_times=1).all().order_by('-add_time')
    user_counts = models.UserProfile.objects.all().count()   # 总人数
    dict1 = {}
    dict2 = {}
    i = 0
    for p in papers:
        testeds = p.record_set.filter(is_valid=True).all()     # 所有成绩
        tested_counts = testeds.count()                        # 已考人数

        # 平均成绩
        score1 = 0
        for s in testeds:
            score1 += s.grade
        avg_score1 = round(score1/(tested_counts if tested_counts else 1))

        # 我的成绩
        rcds = models.Record.objects.filter(userid=request.user, is_valid=True, papername=p).first()    # 我的成绩
        if rcds:
            is_tested = False
        else:
            is_tested = True

        if is_tested:
            dict1[i] = [p, tested_counts]
        else:
            # 已经考试的成绩，只显示最近一个月的，不然太多了
            dt = datetime.datetime.now() - datetime.timedelta(days=30)
            t = p.add_time
            pt = datetime.datetime(year=t.year, month=t.month, day=t.day, hour=t.hour,
                                   second=t.second, minute=t.minute, tzinfo=None)
            if pt > dt:
                dict2[i] = [p, round((tested_counts/user_counts)*100, 2), avg_score1, rcds.grade]
        i += 1
    testing_count = len(dict1)
    tested_count = len(dict2)

    # todo 已完成的考试显示内容太多，可以只显示前十
    return render(request, 'paper_info.html', locals())


# 展示试卷列表
class ListPaperView(View):
    def get(self, request):

        # 试卷信息，选择适合我的试卷,一般用户显示能考试的，管理员显示所有适合的
        querysets1 = models.TestPaper.objects.filter(applyto__pk=request.user.unit.pk).all().order_by('-add_time')  # 按时间降序
        if querysets1 and request.user.getUserRoleName() in ['管理员', '超级管理员']:
            ls = []
            for q in querysets1:
                if q.chance_times != 0:
                    count1 = q.record_set.filter(userid=request.user, is_valid=True).all().count()
                    count = q.chance_times - count1
                    if count == 0:
                        continue
                ls.append(q)
            querysets1 = ls
        count0 = len(querysets1)


        if querysets1:
            querysets = get_pagenator(request, querysets1)
            return render(request, 'paper_list.html', locals())
        else:
            a = mark_safe('系统当前无适合您的试卷!<a href="javascript:history.go(-1);">返回</a>')
            return HttpResponse(a)

# 编辑试卷信息
def editpaper(request, pid):
    paper = TestPaper.objects.get(pk=pid)
    page = request.GET.get('page', 1)
    if request.method == 'GET':
        form_obj = TestPaperModelForm(instance=paper)
        return render(request, 'paper_edit.html', locals())
    else:
        form_obj = TestPaperModelForm(request.POST, instance=paper)
        form_obj.save()
        return redirect('/exam/papers/list/?page={}'.format(page))

# 删除试卷
def deletepaper(request, pid):
    TestPaper.objects.get(pk=pid).delete()
    page = request.GET.get('page', 1)
    return redirect('/exam/papers/list/?page={}'.format(page))



# 展示所有试卷
def list_all_paper(request):
    # 试卷信息
    querysets1 = models.TestPaper.objects.filter().all().order_by('-add_time')  # 按时间降序
    if querysets1:
        querysets = get_pagenator(request, querysets1)
        return render(request, 'paper_list_all.html', locals())
    else:
        a = mark_safe('系统无试卷!<a href="javascript:history.go(-1);">返回</a>')
        return HttpResponse(a)

# 展示试卷成绩
def result_of_paper(request, pid):
    paper = models.TestPaper.objects.get(pk=pid)
    rcds = models.Record.objects.filter(papername=paper).all()
    return render(request, 'result_of_paper.html', locals())

# 导出试卷成绩
def export_result(request, pid):
    paper = models.TestPaper.objects.get(pk=pid)
    rcds = models.Record.objects.filter(papername=paper).all()
    writeToExcel(rcds, paper.paper_name)

    # 返回数据
    filename = 'tmp_results/' + paper.paper_name + '.xls'
    file = open(filename, 'rb')
    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'  # 设置头信息，告诉浏览器这是个文件
    response['Content-Disposition'] = 'attachment;filename="result{}.xls"'.format(
        datetime.datetime.now().strftime('%Y%m%d%H%M%S'))
    # todo 这里文件名不能显示中文，需要设置响应头
    return response

# 将数据写入excel
def writeToExcel(queryset, filename):
    import xlwt
    wb = xlwt.Workbook()
    ws = wb.add_sheet(filename)

    # 写入标题行
    ws.write(0, 0, '序号')
    ws.write(0, 1, '单位')
    ws.write(0, 2, '姓名')
    ws.write(0, 3, '成绩')
    ws.write(0, 4, '有效')
    ws.write(0, 5, '时间')

    # 写入数据
    rowNum = 1
    for q in queryset:
        ws.write(rowNum, 0, rowNum)
        ws.write(rowNum, 1, q.userid.unit.unit_name)
        ws.write(rowNum, 2, q.userid.nick_name)
        ws.write(rowNum, 3, q.grade)
        ws.write(rowNum, 4, '有效' if q.is_valid else '无效')

        # 转换时间的时区，设置为空
        p = q.test_time.strftime("%Y年%m月%d日%H时%M分%S秒")
        ws.write(rowNum, 5, p)
        # dy('写入第{}行成功'.format(rowNum))
        rowNum += 1

    wb.save('tmp_results/'+filename+'.xls')

# 编辑试卷信息-all
def editpaperall(request, pid):
    paper = TestPaper.objects.get(pk=pid)
    page = request.GET.get('page', 1)
    if request.method == 'GET':
        form_obj = TestPaperModelForm(instance=paper)
        return render(request, 'paper_edit_all.html', locals())
    else:
        form_obj = TestPaperModelForm(request.POST, instance=paper)
        form_obj.save()
        return redirect('/exam/papers/list/all/?page={}'.format(page))

# 删除试卷-all
def deletepaperall(request, pid):
    page = request.GET.get('page', 1)
    TestPaper.objects.get(pk=pid).delete()
    return redirect('/exam/papers/list/all/?page={}'.format(page))






# 创建试卷
def create_paper(request):
    if request.method == 'GET':
        # 单位信息，适用人群
        units = Units.objects.all()

        # 题库信息
        qsbks = QuestionInfo.objects.all()
        return render(request, 'paper_create.html', locals())
    else:

        # 试卷名、适用人员列表、题库列表
        pname = request.POST.get('pname', '')
        testtime = request.POST.get('testtime')
        applyto = request.POST.getlist('applyto', '')
        qbs = request.POST.getlist('qb', '')
        is_random = request.POST.get('is_random', None)
        test_count = request.POST.get('test_count', 0)
        if is_random:
            is_random = 1
        else:
            is_random = 0

        # 各型题库数量
        single = request.POST.get('single', 0)
        multi = request.POST.get('multi', 0)
        judge = request.POST.get('judge', 0)
        fillin = request.POST.get('fillin', 0)
        ans = request.POST.get('ans', 0)
        my_choice = [single, multi, judge, fillin, ans]

        # 获取所有题目
        ques_list = [[], [], [], [], []]
        for qb in qbs:
            obj = QuestionInfo.objects.get(id=qb)
            qs = obj.questionbank_set.all()
            for q in qs:
                ques_list[q.question_type].append(q)

        # 随机获取题库信息
        ques_list1 = [[], [], [], [], []]
        for i in range(5):
            lst = list(range(len(ques_list[i])))
            tmpls = shuffle(lst, int(my_choice[i]))
            for j in tmpls:
                ques_list1[i].append(ques_list[i][j])


        return render(request, 'paper_preview.html', {'ques_list1': ques_list1,
                                                      'pname': pname,
                                                      'applyto': json.dumps(applyto),
                                                      'qbs': json.dumps(qbs),
                                                      'testtime': testtime,
                                                      'is_random': is_random,
                                                      'test_count': test_count,
                                                      'my_choice': json.dumps(my_choice),
                                                      })

# 刷新题库数量
def refresh_qb_count(request):
    # 接收前端传过来的试卷编号列表
    ls = json.loads(request.GET.get('paper_ids', ''))
    qblist = [0, 0, 0, 0, 0]
    for i in ls:
        objs = QuestionBank.objects.filter(question_bank_name__pk=i).all()
        # dy(i, objs)
        for item in objs:
            qblist[item.question_type] += 1
    return HttpResponse(json.dumps(qblist))

# 保存创建的试卷
def save_paper(request):

    # 题库、适用人员、题库信息、试卷名
    data = json.loads(request.POST.get('data', ''))
    applyto = json.loads(request.POST.get('applyto', ''))
    my_choice = json.loads(request.POST.get('my_choice', ''))
    is_random = request.POST.get('is_random', 0)
    test_count = request.POST.get('test_count')
    qbs = json.loads(request.POST.get('qbs', ''))

    # 创建试卷
    pname = request.POST.get('pname')
    testtime = request.POST.get('testtime')
    if int(is_random):
        paper = TestPaper.objects.create(
            paper_name=pname, create_person=request.user,
            test_time=testtime, is_random=is_random,
            singel_count=my_choice[0], multi_count=my_choice[1],
            judge_count=my_choice[2], fillin_count=my_choice[3],
            ans_count=my_choice[4], chance_times=test_count
        )
    else:
        paper = TestPaper.objects.create(
            paper_name=pname, create_person=request.user,
            test_time=testtime, is_random=is_random, chance_times=test_count
        )
        for qitme in data:
            obj = QuestionBank.objects.get(pk=qitme)
            obj.pid.add(paper)

    paper.applyto.set(applyto)
    paper.include_qinfo.set(qbs)
    return HttpResponse('sucess')

def get_score(subjects):
    score = 0
    score_all = 0
    right = 0
    wrong = 0
    wronglist = []
    for item in subjects:
        # 去掉非题号键
        if item.isdigit():
            # 根据题号获得题目对象
            q = models.QuestionBank.objects.get(id=item)
            score_all += q.score
            if q.question_type == 0:  # 单选题
                if subjects.get(item) == q.answer.upper():
                    right += 1
                    score += q.score
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 1:  # 多选题
                if ''.join(subjects.getlist(item)) == q.answer.upper():
                    right += 1
                    score += q.score
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 2:  # 判断题
                # 题库答案要包容，人性化
                ans = '正确'
                if q.answer.split() in ['错误', '错误的', '不正确', '不对', '错', '错的', '×', 'x', 'X']:
                    ans = '错误'

                if subjects.get(item) == ans:
                    right += 1
                    score += q.score
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 3:  # 填空题
                # todo 这里要使用分词包
                # todo 查看错题时展示
                pass
            elif q.question_type == 4:  # 问答题
                pass
            else:
                pass

    # 将考试成绩量化成100分
    if score_all == 0:
        score=0
    else:
        score = round(score*100/score_all, 2)
    return right, wrong, wronglist, score


# 申请重考
def retest_paper(request, pid):
    paper = models.TestPaper.objects.get(id=pid)

    # 创建申请记录
    p = models.ProposalRecord.objects.filter(user=request.user, paper=paper, is_handled=False).first()

    a = mark_safe('<a href="javascript:history.back()">返回</a>')
    if p:
        return HttpResponse('您已申请重考，请联系管理员授权！'+a)
    else:
        models.ProposalRecord.objects.create(user=request.user, paper=paper)

        return HttpResponse('申请成功，请耐心等待！'+a)


# 展示所有申请记录
def list_paper_proposal(request):
    proposals = models.ProposalRecord.objects.filter(is_handled=False).all()
    return render(request, 'paper_list_proposal.html', locals())


# 同意申请
def paper_proposal_agree(request, pid):
    # 申请记录处理
    p = models.ProposalRecord.objects.get(pk=pid)
    p.is_handled = True
    p.save()

    # 成绩记录处理
    rcds = models.Record.objects.filter(papername=p.paper, userid=p.user, is_valid=True).all()
    for rc in rcds:
        rc.is_valid = False
        rc.save()
    return redirect('exam:list_paper_proposal')


# 考试
def test_paper(request, pid):
    if request.method == 'GET':
        paper = models.TestPaper.objects.get(id=pid)

        # 如果限定考试次数，查询考试结果是否存在，存在则不能考试,直接显示成绩并提供补考申请窗口，否则考试
        if paper.chance_times > 0:
            rcds = models.Record.objects.filter(papername=paper, userid=request.user, is_valid=True).all()
            if rcds.count() >= paper.chance_times:
                return render(request, 'paper_result_list.html', locals())
            else:  # 允许次数之内，继续往下执行
                pass

        if paper.is_random:
            qinfos = paper.include_qinfo.all()
            ls = [[], [], [], [], []]
            for qinfo in qinfos:
                qs = qinfo.questionbank_set.all()
                for q in qs:
                    ls[q.question_type].append(q)
            subjects = []
            ls_count = [paper.singel_count, paper.multi_count, paper.judge_count, paper.fillin_count, paper.ans_count]
            for i in range(5):
                subjects += shuffle(ls[i], ls_count[i])
        else:
            subjects = paper.questionbank_set.all().order_by('question_type')  # 一对多关系时查询
        return render(request, 'paper_test.html', locals())
    else:
        right, wrong, wronglist, score = get_score(request.POST)

        # 试卷对象，每考试一次热度+1
        qid = request.POST.get('qid')
        qinfo = TestPaper.objects.get(pk=qid)
        qinfo.hot_nums += 1
        qinfo.save()

        train_type = request.POST.get('train_type')
        starttime = request.POST.get('starttime')

        # 保存训练记录
        trcd = models.TrainRecord.objects.create(title=qinfo.paper_name, train_type=train_type,
                                                 user=request.user, starttime=starttime)
        trcd.wrong.set(wronglist)

        # 保存考试成绩
        models.Record.objects.create(userid=request.user, papername=qinfo, grade=score)
        return redirect('exam:list_result')




# 保存试卷为word文档
def save_paper_word(request, pid):
    """
    1、将试卷生成后写入后台服务器文件夹
    2、根据试卷是随机与否写入
    3、如果是固定出题，直接从文件夹下载传回前端即可
    :param request:
    :param pid:
    :return:
    """
    paper = models.TestPaper.objects.get(pk=pid)

    # 固定模式
    if paper.is_random == 0:
        # 保存试卷到缓存
        papers = paper.questionbank_set.all().order_by('question_type')
        # paperSaveAsWord(papers, paper)

    else:   # 随机模式
        qinfos = paper.include_qinfo.all()
        ls = [[], [], [], [], []]
        for qinfo in qinfos:
            qs = qinfo.questionbank_set.all()
            for q in qs:
                ls[q.question_type].append(q)
        papers = []
        ls_count = [paper.singel_count, paper.multi_count, paper.judge_count, paper.fillin_count, paper.ans_count]
        for i in range(5):
            papers += shuffle(ls[i], ls_count[i])

    paperSaveAsWord(papers, paper)

    # 读取数据，返回
    pname = str(paper.pk) + paper.paper_name
    filename = 'tmp_papers/' + pname + '.docx'
    file = open(filename, 'rb')
    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'  # 设置头信息，告诉浏览器这是个文件
    response['Content-Disposition'] = 'attachment;filename="{}.docx"'.format(
        datetime.datetime.now().strftime('%Y%m%d%H%M%S'))
    # todo 这里文件名不能显示中文，需要设置响应头
    return response


# 将数据集保存为word文档
def paperSaveAsWord(queryset, paper):
    # 文件名
    pname = str(paper.pk)+paper.paper_name
    filename = 'tmp_papers/'+pname+'.docx'

    from docx import Document
    doc = Document()
    for q in queryset:
        # todo 这里适用可以不要
        str1 = '[题型:{}][答案:{}][分数:{}][难度:{}][适用:所有人]'.format(q.get_question_type_display(),
                                                             q.answer,
                                                             q.score,
                                                             q.get_difficulty_display(),)
        doc.add_paragraph(str1)
        doc.add_paragraph(q.title)
        doc.add_paragraph()

    # 如果文件缓存目录不存在，新建
    if not os.path.exists('tmp_papers'):
        os.mkdir('tmp_papers')
    doc.save(filename)





# 开始考试
class StartExamView(View):
    def get(self, request, pid):
        paper = models.TestPaper.objects.filter(id=pid).first()
        subjects = paper.questionbank_set.all()  # 多对多关系时查询
        return render(request, 'start_exam.html', locals())

    def post(self, request, pid):
        dy(pid, request.POST)

        # 计算得分, 存储错题
        paper = models.TestPaper.objects.filter(id=pid).first()
        scores = self.calculate(request.POST, request, paper)

        # 存入结果到考试记录中
        obj = models.Record.objects.create(userid=request.user, papername=paper, grade=scores)
        obj.save()

        # todo 这里需要重定向
        return redirect('exam:list_result')

    def calculate(self, subjects, request, paper):
        sum_score = 0
        for item in subjects:

            # 去掉非题号键
            if item.isdigit():

                # 根据题号获得题目对象
                q = models.QuestionBank.objects.get(id=item)

                # 查询用户当前错题存在否
                wr = models.WrongRecord.objects.filter(user=request.user, question=q).first()  # 错题
                if q.question_type == 0:  # 单选题
                    if subjects.get(item) == q.answer.upper():
                        # 正确，得分
                        sum_score += q.score

                        # 错误次数减1
                        if wr:
                            wr.error_times -= 1 if wr.error_times > 0 else 0  # 可以为负，表示已经做对次数
                            wr.save()
                    else:
                        # 错误不得分，保存错题。已存在，增加次数，不存在新建。
                        if wr:
                            wr.error_times += 1
                            wr.my_ans = subjects.get(item)
                            wr.save()
                        else:
                            models.WrongRecord.objects.create(user=request.user, question=q, error_times=1,
                                                              paper=paper, my_ans=subjects.get(item))  # 错题

                elif q.question_type == 1:  # 多选题
                    if ''.join(subjects.getlist(item)) == q.answer.upper():
                        sum_score += q.score
                        if wr:
                            wr.error_times -= 1 if wr.error_times > 0 else 0
                            wr.save()
                    else:
                        if wr:
                            wr.error_times += 1
                            wr.save()
                        else:
                            models.WrongRecord.objects.create(user=request.user, question=q, error_times=1,
                                                              paper=paper)  # 错题
                elif q.question_type == 2:  # 判断题
                    # 题库答案要包容，人性化
                    ans = '正确'
                    if q.answer.split() in ['错误', '错误的', '不正确', '不对', '错', '错的', '×', 'x', 'X']:
                        ans = '错误'

                    if subjects.get(item) == ans:
                        sum_score += q.score
                        if wr:
                            wr.error_times -= 1 if wr.error_times > 0 else 0
                            wr.save()
                    else:
                        if wr:
                            wr.error_times += 1
                            wr.save()
                        else:
                            models.WrongRecord.objects.create(user=request.user, question=q, error_times=1,
                                                              paper=paper)  # 错题
                elif q.question_type == 3:  # 填空题
                    # todo 这里要使用分词包
                    # todo 查看错题时展示
                    pass
                elif q.question_type == 4:  # 问答题
                    pass
                else:
                    pass
        return sum_score

# 成绩列表
class ListResultView(View):
    def get(self, request):
        # 一般用户显示自己的成绩，管理员显示所有人成绩
        if request.user.getUserRoleName() == '一般用户':
            # 只显示自己的考试记录
            grade = models.Record.objects.filter(userid=request.user).all().order_by('-test_time')
        else:
            grade = models.Record.objects.all().order_by('-test_time')
        grade = get_pagenator(request, grade)
        return render(request, 'list_result.html', {'grade': grade})


def delete_result(request, rid):
    page = request.GET.get('page', 1)
    r = models.Record.objects.get(pk=rid).delete()
    return redirect('/exam/results/list/?page={}'.format(page))


class WrongRecordView(View):
    def get(self, request):
        wr = models.WrongRecord.objects.filter(user=request.user).all()
        wr = self.get_wrongs(wr)
        return render(request, 'wrong-record.html', {'wr': wr})

    def get_wrongs(self, wr):
        wdict = {}
        for item in wr:
            paper = item.paper
            if paper in wdict:
                wdict[paper] += 1
            else:
                wdict[paper] = 1
        return wdict


class ListWrongsView(View):
    def get(self, request):

        # 查看错题详情
        pid = request.GET.get('pid', None)
        if pid:
            objs = models.WrongRecord.objects.filter(paper__id=pid).all()
            return render(request, 'start_wrongs.html', {'objs': objs})

        # 试卷信息
        querysets = models.TestPaper.objects.all().order_by('-add_time')  # 按时间降序
        papers_list = []
        for paper in querysets:
            wrongs = paper.wrongrecord_set.filter(user=request.user, error_times__gt=0).all()
            dy(wrongs)
            if wrongs:
                papers_list.append(paper)
        if papers_list:
            querysets = get_pagenator(request, papers_list)
            return render(request, 'list_wrongs.html', {'querysets': querysets})
        else:
            return HttpResponse('当前无错题！')


# 展示题库中心信息
def qsbk_info(request):
    return render(request, 'qsbk_info.html')

# 题库信息列表
class ListQSBKView(View):
    def get(self, request):
        if request.user.getUserRoleName() == '一般用户':
            querysets1 = models.QuestionInfo.objects.filter(create_person=request.user).all().order_by('-add_time')
        else:
            querysets1 = models.QuestionInfo.objects.all().order_by('-add_time')
        querysets = get_pagenator(request, querysets1)
        return render(request, 'qsbk_list.html', locals())


# 题库详细信息可编辑每一题
class QuestionBankDetailView(View):
    def get(self, request, qb_id):
        page = request.GET.get('page')
        # 编辑
        edit = request.GET.get('edit')
        if edit:
            obj = models.QuestionBank.objects.filter(id=int(edit)).first()
            form_obj = examforms.QuestionModelForm(instance=obj)

            return render(request, 'questionbank_edit.html', {'form_obj': form_obj, 'qid': qb_id, 'page': page})

        # 删除
        delete = request.GET.get('del')
        if delete:
            models.QuestionBank.objects.get(id=delete).delete()

        qb_info = models.QuestionInfo.objects.get(id=qb_id)
        querysets = get_pagenator(request, qb_info.questionbank_set.all(), num=10)
        return render(request, 'questionbank_detail.html', locals())

    def post(self, request, qb_id):
        qid = request.GET.get('qid')
        obj = models.QuestionBank.objects.get(id=qid)
        form_obj = examforms.QuestionModelForm(request.POST, instance=obj)
        if form_obj.is_valid():
            form_obj.save()
            page = request.GET.get('page')
            return redirect('/exam/qsbk/list/{}/detail/?page={}'.format(qb_id, page))
        return render(request, 'questionbank_edit.html', {'form_obj': form_obj})


# 删除题库
def qsbk_delete(request, qb_id):
    page = request.GET.get('page', 1)
    QuestionBank.objects.filter(question_bank_name__pk=qb_id).delete()
    QuestionInfo.objects.filter(pk=qb_id).delete()
    return redirect('/exam/qsbk/list/?page={}'.format(page))


# 编辑题库
def qsbk_edit(request, qb_id):
    if request.method == 'GET':
        obj = QuestionInfo.objects.filter(pk=qb_id).first()
        form_obj = QuestionInfoModelForm(instance=obj)

        return render(request, 'qsbk_edit.html', locals())
    else:
        # dy('post')
        page = request.GET.get('page', 1)
        obj = QuestionInfo.objects.filter(pk=qb_id).first()
        QuestionInfoModelForm(request.POST, instance=obj).save()
        return redirect('/exam/qsbk/list/?page={}'.format(page))

# 导入题库
def qsbk_import(request):
    if request.method == 'GET':
        units = models.Units.objects.all()
        return render(request, 'qsbk_import.html', locals())
    else:
        # 保存word文档
        f = request.FILES.get('filename', '')
        if f:

            # 如果文件缓存目录不存在，新建
            if not os.path.exists('tmpfiles'):
                os.mkdir('tmpfiles')

            # 写入文件
            with open('tmpfiles/' + f.name, 'wb+') as destination:
                for chunk in f.chunks():
                    destination.write(chunk)
            return HttpResponse(json.dumps({'msg': '上传成功！', "status": 'success'}), content_type="application/json")
        else:
            return JsonResponse({'msg': '上传失败！', "status": 'fail'})

# 下载题库模板
def qsbk_download(request):
    types = request.GET.get('type')
    a = types+'_template.'
    if types == 'word':
        a += 'docx'
    if types == 'excel':
        a += 'xls'

    # dy('下载题库模板', a)
    file = open('static/' + a, 'rb')
    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'  # 设置头信息，告诉浏览器这是个文件
    response['Content-Disposition'] = 'attachment;filename="{}"'.format(a)
    return response

from utils.myexcel import readExcelToData
# 导入题库预览
def qsbk_show(request):

    fname = request.FILES.get('filename').name
    qbname = request.POST.get('qsbk_name')
    applyto = json.dumps(request.POST.getlist('applyto', ''))
    # dy('题库预览时回传数据', fname, qbname, applyto)

    data = {}
    types = request.GET.get('type')
    filepath = 'tmpfiles/' + fname
    if types == 'word':
        try:
            data = QuestionBankClass(filepath).get_questionbank()
        except:
            pass
    else:
        data = readExcelToData(filepath)
        dy(data)

    return render(request, 'qsbk_show.html', locals())

# 题库保存
def qsbk_save(request):
    qbname = request.POST.get('qbname')
    applyto = json.loads(request.POST.get('applyto', ''))
    q_obj = QuestionInfo.objects.create(name=qbname, create_person=request.user)
    q_obj.applyto.set(applyto)
    # todo 可以做适用人员

    data = json.loads(request.POST.get('data', ''))
    for d in data:
        tmpd = {}
        for key, val in d.items():
            tmpd[key] = val
        tmpd['question_bank_name'] = q_obj
        # dy('保存字典', tmpd)
        QuestionBank.objects.create(**tmpd)
    return redirect('exam:list_questionbank_info')





# 训练中心
def train_info(request):
    # 训练记录，总时长、今年以来时长
    trcds = models.TrainRecord.objects.filter(user=request.user).all()
    time_all = 0
    time_thisyear = 0
    for tr in trcds:
        time_all += tr.get_time_diff()
        time_thisyear += tr.get_time_of_this_year()

    time_all = math.ceil(time_all/3600)    # 转换成小时
    time_thisyear = math.ceil(time_thisyear/60)   # 转换成分钟

    # 题库总数，题目数量
    qbinfos = models.QuestionInfo.objects.all()
    qbcounts = 0
    for qbf in qbinfos:
        qbcounts += qbf.questionbank_set.count()

    # 题库使用率，试卷所用数量占题库数量比值
    qb_set = set()
    for p in models.TestPaper.objects.all():
        p1 = p.questionbank_set.all()
        for p2 in p1:
            qb_set.add(p2)
    if qbcounts != 0:
        used_rate = round((len(qb_set) / qbcounts)*100)
    else:
        used_rate = 0

    # 所有适合我的题目
    my_qs_count = 0
    qsinfo_mine = models.Units.objects.get(pk=request.user.unit.pk).questioninfo_set.all()
    for qsinfo in qsinfo_mine:
        my_qs_count += qsinfo.questionbank_set.count()
    return render(request, 'train_info.html', locals())

# 训练列表，显示所有可以练习的题库
def train_list(request):

    # 获取当前用户所在单位适用的所有题库信息
    qb_infos = request.user.unit.questioninfo_set.all().order_by('-add_time')
    return render(request, 'train_list.html', locals())


# 随机100题练习
def random_test(request, qid):
    # 题库信息
    qb = QuestionInfo.objects.get(pk=qid)

    # 题库对应的题目集
    qs = qb.questionbank_set.all().order_by('question_type')

    # 如果数量不足100，直接返回全部，大于100，随机挑选100题
    if qs.count() <= 100:
        # dy('返回全部！')
        return render(request, 'train_random100.html', locals())
    else:
        # dy('随机返回100题！')
        qs = shuffle(list(qs), 100)
        return render(request, 'train_random100.html', locals())

# 顺序练习
def train_sequence(request, qid):
    qb = QuestionInfo.objects.get(pk=qid)
    qs = qb.questionbank_set.all().order_by('question_type')
    return render(request, 'train_sequence.html', locals())



def calculate(subjects):
    right = 0
    wrong = 0
    wronglist = []
    for item in subjects:
        # 去掉非题号键
        if item.isdigit():
            # 根据题号获得题目对象
            q = models.QuestionBank.objects.get(id=item)
            if q.question_type == 0:  # 单选题
                if subjects.get(item) == q.answer.upper():
                    right += 1
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 1:  # 多选题
                if ''.join(subjects.getlist(item)) == q.answer.upper():
                    right += 1
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 2:  # 判断题
                # 题库答案要包容，人性化
                ans = '正确'
                if q.answer.split() in ['错误', '错误的', '不正确', '不对', '错', '错的', '×', 'x', 'X']:
                    ans = '错误'

                if subjects.get(item) == ans:
                    right += 1
                else:
                    wronglist.append(q)
                    wrong += 1
            elif q.question_type == 3:  # 填空题
                # todo 这里要使用分词包
                # todo 查看错题时展示
                pass
            elif q.question_type == 4:  # 问答题
                pass
            else:
                pass
    return right, wrong, wronglist

# 校对模式
def train_check(request, qid):
    if request.method == 'GET':
        qb = QuestionInfo.objects.get(pk=qid)
        qs = qb.questionbank_set.all().order_by('question_type')
        return render(request, 'train_check.html', locals())
    else:
        dy(request.POST)
        # 通过判断start_time是否存在，确定为ajax刷新还是提交
        start_time = request.POST.get('start_time', '')
        if start_time:
            # 训练方式
            train_type = request.POST.get('train_type', '')
            tt = '参加【' + train_type + '】练习'

            qid = request.POST.get('qid', '')
            qb = QuestionInfo.objects.get(pk=qid)
            # dy('训练情况', qid, train_type, start_time)
            time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            obj = models.TrainRecord.objects.create(title=qb.name, train_type=tt,
                                                    starttime=start_time, user=request.user,
                                                    time=time)
            # dy(obj)
            return redirect('exam:train_list')
        else:
            quesid = request.POST.get('name')
            val = request.POST.get('val')
            qs = QuestionBank.objects.get(pk=quesid)
            if qs.answer == val:
                return HttpResponse('right')
            else:
                return HttpResponse('wrong')

# 提交练习
def train_submit(request):
    # dy(request.POST)
    right, wrong, wronglist = calculate(request.POST)
    qid = request.POST.get('qid')
    qinfo = QuestionInfo.objects.get(pk=qid)
    train_type = '参加【' + request.POST.get('train_type') + '】练习'
    starttime = request.POST.get('starttime')

    # 结束时间
    time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # todo 这里如果没答题将没有结果
    trcd = models.TrainRecord.objects.create(title=qinfo, train_type=train_type,
                                             user=request.user, time=time, starttime=starttime)
    trcd.wrong.set(wronglist)
    return redirect('exam:train_record')

# 训练记录
def train_record(request):
    rcds1 = models.TrainRecord.objects.filter(user=request.user).all().order_by('-time')
    if rcds1:
        rcds = get_pagenator(request, rcds1)
        return render(request, 'train_record.html', locals())
    else:
        return HttpResponse('当前无记录！')

# 训练记录
def train_allrecord(request):
    rcds = models.TrainRecord.objects.all().order_by('-time')
    if rcds:
        rcds = get_pagenator(request, rcds)
        return render(request, 'train_allrecord.html', locals())
    else:
        return HttpResponse('当前无记录！')

# 删除训练记录
def delete_record(request, wid):
    models.TrainRecord.objects.get(pk=wid).delete()
    page = request.GET.get('page', 1)
    isall = request.GET.get('all')
    if isall == 'yes':
        return redirect('/exam/train/allrecord/?page={}'.format(page))
    else:
        return redirect('/exam/train/record/?page={}'.format(page))

# 显示训练记录中的错题
def listwrong(request, wid):
    if request.method == 'GET':
        trcd = models.TrainRecord.objects.get(pk=wid)
        trcds = trcd.wrong.select_related().all().order_by('question_type')
        return render(request, 'train_listwrongs.html', locals())
    else:
        quesid = request.POST.get('name')
        val = request.POST.get('val')
        qs = QuestionBank.objects.get(pk=quesid)
        dy(qs.answer, val)
        if qs.answer == val:
            return JsonResponse({'status': 'right', 'ans': qs.answer})
        else:
            return JsonResponse({'status': 'wrong', 'ans': qs.answer})

# 问卷调查首页
def questionnaire_info(request):
    return render(request, 'questionnaire_info.html')

