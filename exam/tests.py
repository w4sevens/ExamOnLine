from django.test import TestCase
from docx import Document
import os
from utils.utils import dy
queryset = ['shshdlkfsjld', 'shdklfsjkdjf']
def paperSaveAsWord(queryset):
    doc = Document()
    doc.add_paragraph('段落1')
    doc.add_paragraph('段落2')
    doc.add_paragraph('段落3')

    # 如果文件缓存目录不存在，新建
    if not os.path.exists('tmp_papers'):
        os.mkdir('tmp_papers')
    doc.save('../tmp_papers/good1.docx')

import xlwt, xlrd
def readExcelToData(filepath):
    dy('进入')
    data = []
    file = xlrd.open_workbook(filepath)
    dy(type(file))
    table = file.sheet_by_index(0)
    print(type(table))

if __name__ == '__main__':
    # readExcelToData('../tmpfiles/excel_template.xlsx')
    dy('good')